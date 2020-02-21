# -*- coding: utf-8 -*-
"""
Created on Wen Jan 09 2019
@author: Rt-Rakesh

This script is a utility generate a eda document of the csv.
Usage: Should be used by calling in jupyter notbooks.
"""
import os
import glob
import pandas as pd
import xml.etree.ElementTree as ET
import numpy as np
from docx import Document


def _generate_count_df(data_df):
    count_df = data_df.groupby('label').agg({"label": "count", "path": "nunique"})
    count_df.columns = ['Total Count', 'Unique Document']
    count_df.index.names = ['Labels']
    count_df.sort_values('Total Count', ascending=False, inplace=True)
    return count_df


def generate_EDA(data_file_path,
                 odapi_path,
                 label_map_path,
                 save_path,
                 detection_tf_record_path,
                 iou_threshold=0.5,
                 confidence_threshold=0.5,
                 display_cm=False):
    """
    This Function generates the confusion matrix, presicion recall document for all the labels handles by the moadel.
    -----
    Args:
    -----
    1.odapi_path: --str The object detection api source code path( This is optional)
    2.label_map_path: --str  The path to the label map.
    3.detection_tf_record_path: -- str The path to the detections tf record path.
    4.iou_threshold: --float Deault:0.5
    5.confidence_threshold: --float Default:0.5
    6.display_cm: --Boolean Deafault:False Displays the confusion matrix, the label wise presicion and recall.
    ---------
    Returns:
    ---------
    Saves confusion matrix image and precision recall document.
    """

    data_df = pd.read_csv(data_file_path)
    labels = data_df['label'].unique().tolist()
    total_unique_images = data_df['path'].nunique()

    count_df = _generate_count_df(data_df)

    width_list = []
    height_list = []
    xmin = data_df['xmin'].tolist()
    xmax = data_df['xmax'].tolist()
    for x, y in zip(xmin, xmax):
        width = int(y)-int(x)
        width_list.append(width)
    ymin = data_df['ymin'].tolist()
    ymax = data_df['ymax'].tolist()
    for x, y in zip(ymin, ymax):
        height = int(y)-int(x)
        height_list.append(height)

    data_df['label_width'] = width_list
    data_df['label_height'] = height_list

    size_df = data_df.groupby(['label']).agg({"label_width": [min, np.median, max, ], "label_height": [min, np.median, max]})
    size_df.columns.set_levels(['Min', 'Median', 'Max'], level=1, inplace=True)
    size_df.columns.set_levels(['Lable Width', 'Label Height'], level=0, inplace=True)
    size_df.index.names = ['Labels']

    size_df.reset_index(level=0, inplace=True)
    count_df.reset_index(level=0, inplace=True)

    doc = docx.Document()
    document.add_heading('Data Analysis ', 0)
    p = doc.add_paragraph()
    p.add_run('The total count for labels and the unique documents').bold = True
    p.add_run('')
    p.add_run('')
    t = doc.add_table(count_df.shape[0]+1, count_df.shape[1])
    for j in range(count_df.shape[-1]):
        t.cell(0, j).text = count_df.columns[j]
    for i in range(count_df.shape[0]):
        for j in range(count_df.shape[-1]):
            t.cell(i+1, j).text = str(count_df.values[i, j])

    p = doc.add_paragraph()
    p.add_run('The Width and Height analysis of the annotations').bold = True
    p.add_run('')
    p.add_run('')
    t = doc.add_table(size_df.shape[0]+1, size_df.shape[1])
    for j in range(size_df.shape[-1]):
        t.cell(0, j).text = size_df.columns[j]
    for i in range(size_df.shape[0]):
        for j in range(size_df.shape[-1]):
            t.cell(i+1, j).text = str(size_df.values[i, j])
    doc.save('./test.docx')
